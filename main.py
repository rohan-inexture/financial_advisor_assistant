from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.prompts import ChatPromptTemplate
from langchain_community.document_loaders import PyPDFLoader
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain.schema import Document as LangchainDocument
from langchain_community.vectorstores import Chroma
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.chains import create_retrieval_chain
from dotenv import load_dotenv
from docx import Document
import streamlit as st
import pandas as pd
import os
import json
import plotly.graph_objects as go
import plotly.express as px
import shutil
import time

load_dotenv()


def main_document():
    # Set page configuration
    st.set_page_config(
        page_title="Financial Assistant",
        page_icon="img/Yellow Favicon.png",
        layout="wide"
    )

    # Initialize session state for data
    if "data" not in st.session_state:
        st.session_state.data = None
    if "df" not in st.session_state:
        st.session_state.df = None
    if "text_content" not in st.session_state:
        st.session_state.text_content = None
    if "file_processed" not in st.session_state:
        st.session_state.file_processed = False
    if "analysis_results" not in st.session_state:
        st.session_state.analysis_results = None
    if "chart_config" not in st.session_state:
        st.session_state.chart_config = []
    if "raw_text" not in st.session_state:
        st.session_state.raw_text = ""
    if "conversation_history" not in st.session_state:
        st.session_state.conversation_history = []
    if "allocations" not in st.session_state:
        st.session_state.allocations = {}
    if "investment_types" not in st.session_state:
        st.session_state.investment_types = []
    if "show_allocation" not in st.session_state:
        st.session_state.show_allocation = False


    # Add the logo to the sidebar
    st.sidebar.image("img/image.png", use_container_width=True)

    st.title("Financial Assistant 💰")

    llm = ChatGoogleGenerativeAI(model="gemini-2.0-flash")
    # llm = ChatOpenAI(model_name="gpt-4o-mini-2024-07-18", temperature=0.7)

    # Enhanced prompt for financial context
    prompt_temp = ChatPromptTemplate.from_template(
        """
        You are a specialized financial assistant. Answer the user's questions based on the provided context and their financial data.
        If the user requests budget adjustments or savings projections, use the financial data in your context.

        <context>
        {context}
        </context>

        User query: {input}

        When responding:
        1. If answering a question about documents, provide specific information from them
        2. If asked about budgets, refer to the current budget data
        3. If requested to make financial projections, show calculations clearly
        4. If asked to adjust a budget, provide the updated values and rationale
        5. Be precise with numbers and financial advice
        6. If you cannot fulfill a request from the context, politely explain why

        Your response should be helpful, accurate, and tailored to financial matters.
        """
    )

    # Create a temporary directory to save uploaded files
    temp_dir = "temp_files"
    if not os.path.exists(temp_dir):
        os.makedirs(temp_dir)

    def load_docx(file_path):
        """Function to load and extract text from a DOCX file"""
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return '\n'.join(full_text)

    def load_csv(file_path):
        """Function to load and extract text from a CSV file"""
        try:
            df = pd.read_csv(file_path)
            return df.to_string()
        except Exception as e:
            return f"Error loading CSV: {str(e)}"

    def vector_embedding(uploaded_file):
        df = None
        if "vectors" not in st.session_state:
            st.session_state.embeddings = GoogleGenerativeAIEmbeddings(model="models/text-embedding-004")
            # st.session_state.embeddings = OpenAIEmbeddings(model="text-embedding-3-small")  # or "text-embedding-ada-002"

            documents = []

            # Save the uploaded file temporarily
            temp_file_path = os.path.join(temp_dir, uploaded_file.name)
            with open(temp_file_path, "wb") as f:
                f.write(uploaded_file.getbuffer())

            # Load the file content based on the file type
            if uploaded_file.name.endswith('.pdf'):
                loader = PyPDFLoader(temp_file_path)
                loaded_docs = loader.load()
                for doc in loaded_docs:
                    documents.append(LangchainDocument(page_content=doc.page_content, metadata=doc.metadata))
            elif uploaded_file.name.endswith('.docx'):
                docx_text = load_docx(temp_file_path)
                documents.append(LangchainDocument(page_content=docx_text, metadata={'source': temp_file_path}))
            elif uploaded_file.name.endswith('.csv'):
                csv_text = load_csv(temp_file_path)
                # Also load into financial data if it seems like financial data
                try:
                    df = pd.read_csv(temp_file_path)
                except:
                    pass
                documents.append(LangchainDocument(page_content=csv_text, metadata={'source': temp_file_path}))

            st.session_state.text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
            st.session_state.final_documents = st.session_state.text_splitter.split_documents(documents)
            st.session_state.vectors = Chroma.from_documents(st.session_state.final_documents,
                                                             st.session_state.embeddings)

            # Delete the temporary files after vector creation
            temp_file_path = os.path.join(temp_dir, uploaded_file.name)
            os.remove(temp_file_path)
            return df, documents

    def detect_date_columns(df):
        """Detect columns that might contain dates"""
        date_columns = []

        for col in df.columns:
            # Skip if column is numeric
            if pd.api.types.is_numeric_dtype(df[col]):
                continue

            # Try to convert to datetime
            try:
                pd.to_datetime(df[col], errors='raise')
                date_columns.append(col)
            except:
                # Check if column name suggests a date
                if any(date_term in col.lower() for date_term in ['date', 'year', 'month', 'day', 'time']):
                    date_columns.append(col)

        return date_columns

    def analyze_data_with_llm(df, text_content=""):
        """Use LLM to analyze data and suggest visualizations"""
        # First try Groq, fall back to Gemini if needed
        try:
            column_info = ", ".join([f"{col} ({df[col].dtype})" for col in df.columns])
            sample_data = df.head(5).to_string()

            prompt = f"""
                You are a financial data analysis agent.

                Analyze the financial data provided below and return structured JSON output with the following keys:

                ```json
                {{
                "summary": "Brief overview of what the data contains",
                "columns": {{
                    "numeric": ["list", "of", "numeric", "columns"],
                    "categorical": ["list", "of", "categorical", "columns"],
                    "temporal": ["list", "of", "date", "columns"]
                }},
                "suggested_charts": [
                    {{
                    "type": "line/bar/scatter/pie/box/histogram",
                    "title": "Suggested chart title",
                    "x_axis": "column_name",
                    "y_axis": "column_name",
                    "color": "optional_column_for_color",
                    "description": "Why this visualization is appropriate"
                    }}
                ],
                "insights": [
                    "key insight 1",
                    "key insight 2"
                ]
                }}

                Document Data:
                Columns: {column_info} 
                Sample data:{sample_data}

            """

            llm_response = llm.invoke(prompt)
            response = llm_response.content

            json_content = response.replace("```json", "").replace("```", "").strip()
            return json.loads(json_content)

        except Exception as e:
            st.error(f"Error analyzing data with LLM: {e}")

            return None

    def create_visualization(chart_config, df):
        """Create a Plotly visualization based on the chart configuration"""
        chart_type = chart_config.get("type", "").lower()
        title = chart_config.get("title", "Chart")
        x_axis = chart_config.get("x_axis")
        y_axis = chart_config.get("y_axis")
        color = chart_config.get("color")

        if not x_axis and not y_axis:
            return None

        try:
            if chart_type == "line":
                # Convert to datetime if x_axis is a date column
                if x_axis in detect_date_columns(df):
                    df[x_axis] = pd.to_datetime(df[x_axis], errors='coerce')

                if color:
                    fig = px.line(df, x=x_axis, y=y_axis, color=color, title=title)
                else:
                    fig = px.line(df, x=x_axis, y=y_axis, title=title)

            elif chart_type == "bar":
                if color:
                    fig = px.bar(df, x=x_axis, y=y_axis, color=color, title=title)
                else:
                    fig = px.bar(df, x=x_axis, y=y_axis, title=title)

            elif chart_type == "scatter":
                if color:
                    fig = px.scatter(df, x=x_axis, y=y_axis, color=color, title=title)
                else:
                    fig = px.scatter(df, x=x_axis, y=y_axis, title=title)

            elif chart_type == "pie":
                fig = px.pie(df, names=x_axis, values=y_axis, title=title)

            elif chart_type == "box":
                fig = px.box(df, x=x_axis, y=y_axis, color=color, title=title)

            elif chart_type == "histogram":
                fig = px.histogram(df, x=x_axis, title=title)

            else:
                # Default to a table view
                fig = go.Figure(data=[go.Table(
                    header=dict(values=list(df.columns)),
                    cells=dict(values=[df[col] for col in df.columns])
                )])
                fig.update_layout(title=title)

            # Update layout for better appearance
            fig.update_layout(
                title=title,
                template="plotly_white",
                margin=dict(l=20, r=20, t=40, b=20),
                height=450
            )

            return fig

        except Exception as e:
            st.error(f"Error creating visualization: {e}")
            return None

    def financial_adviser(preferences_data):
        advisor_prompt = f"""
            You are a certified financial advisor. Analyze the user's financial preferences and develop a comprehensive and personalized financial plan.

            User Financial Preferences: {preferences_data}

            Instructions:

                1) Interpret the user's investment intent and preferences:
                    If an investment horizon is specified (e.g., short/medium/long term), tailor the plan accordingly. If not, suggest a suitable horizon based on general goals and medium risk tolerance.
                    If investment types are provided (e.g., SIPs, FDs, stocks), use them in your recommendations. If missing, suggest suitable instruments based on user profile.
                    If allocation is defined (e.g., 40% SIP, 60% FD), assess its suitability. If not, propose a strategic allocation plan across asset classes.

                2) Develop a clear, actionable investment strategy:
                    Recommend how to deploy the investment amount to meet the stated or inferred goals.
                    Suggest instruments across risk categories (e.g., equity mutual funds, FDs, NPS, PPF, index funds, gold).
                    Provide a proposed portfolio breakdown (percentages and ₹ amounts).
                
                3) Include long-term planning suggestions:
                    Building an emergency fund
                    Retirement planning with moderate risk
                    Provisions for dependent parents if applicable (e.g., healthcare corpus, monthly support)
                    
                4) Offer additional recommendations:
                    Tax-saving strategies (ELSS, PPF, NPS, 80C)
                    Basic insurance needs (life, health)
                    Monthly savings discipline and portfolio rebalancing tips
            
            Important Notes:
                Do not assume or comment on income, expenses, or budgeting unless provided.
                Focus only on making the best use of the investment amount and preferences.
                Use ₹ values where helpful. Be clear, actionable, and goal-oriented.

        """
        
        llm_response = llm.invoke(advisor_prompt)
        response = llm_response.content
        return response

    def investment_strategy_visulizations(response):
        prompt = f"""
            You are a data visualization assistant. Based on the following financial planning response, extract key data points and recommend suitable charts to visualize the investment strategy, allocations, and planning elements.

            Your output must be a JSON object in the format:
            {{
                "suggested_charts": [
                    {{
                    "type": "line/bar/scatter/pie/box/histogram",
                    "title": "Suggested chart title",
                    "x_axis": data,
                    "y_axis": data,
                    "color": data,
                    "description": "Why this visualization is appropriate"
                    }}
                ]
            }}

            Use appropriate chart types to represent allocation (e.g., pie for proportions, bar for category comparison). Do not invent data—only use insights explicitly or implicitly mentioned in the response.

            Here is the financial response: {response}
        """

        llm_response = llm.invoke(prompt)
        response = llm_response.content

        json_content = response.replace("```json", "").replace("```", "").strip()
        return json.loads(json_content)
   
    def plot_suggested_charts(chart_data):
        # Define a color palette
        color_palette = ['#1f77b4', '#ff7f0e', '#2ca02c', '#d62728', '#9467bd', 
                        '#8c564b', '#e377c2', '#7f7f7f', '#bcbd22', '#17becf']
        
        for chart in chart_data["suggested_charts"]:
            try:
                chart_type = chart["type"]
                title = chart["title"]
                
                # Validate and convert x_axis and y_axis data
                if isinstance(chart["x_axis"], str):
                    x = [chart["x_axis"]]  # Convert single string to list
                elif isinstance(chart["x_axis"], (list, tuple)):
                    x = list(chart["x_axis"])
                else:
                    raise ValueError(f"Invalid x_axis data type for chart '{title}'")
                    
                if isinstance(chart["y_axis"], str):
                    y = [float(chart["y_axis"])]  # Convert single string to float
                elif isinstance(chart["y_axis"], (list, tuple)):
                    y = [float(val) if isinstance(val, str) else val for val in chart["y_axis"]]
                else:
                    raise ValueError(f"Invalid y_axis data type for chart '{title}'")
                    
                # Ensure x and y have the same length
                if len(x) != len(y):
                    raise ValueError(f"Mismatched data lengths in chart '{title}': {len(x)} labels vs {len(y)} values")
                
                description = chart.get("description", "")

                if chart_type == "pie":
                    fig = go.Figure(data=[go.Pie(
                        labels=x, 
                        values=y, 
                        marker=dict(colors=color_palette[:len(x)])
                    )])
                    fig.update_layout(
                        title=title,
                        showlegend=True
                    )
                
                elif chart_type == "bar":
                    fig = go.Figure(data=[go.Bar(
                        x=x, 
                        y=y, 
                        marker=dict(color=color_palette[:len(x)])
                    )])
                    fig.update_layout(
                        title=title,
                        xaxis_title="Categories",
                        yaxis_title="Values",
                        showlegend=False
                    )

                else:
                    continue

                # Update common layout properties
                fig.update_layout(
                    title_x=0.5,
                    margin=dict(t=50, b=50, l=50, r=50),
                    height=400
                )

                return fig
                
            except (ValueError, TypeError) as e:
                raise ValueError(f"Data validation error in chart '{title}': {str(e)}")
            except Exception as e:
                raise Exception(f"Unexpected error in chart '{title}': {str(e)}")

    # Create tabs for different functionalities
    tab1, tab2, tab3 = st.tabs(["Chat", "Dashboard", "Advisor"])

    with tab3:
        st.markdown("""
            <style>
            .financial-form {
                background-color: #f8f9fa;
                padding: 20px;
                border-radius: 10px;
                box-shadow: 0 2px 4px rgba(0,0,0,0.1);
            }
            .advice-section {
                margin-top: 20px;
                padding: 20px;
                border-radius: 10px;
                background-color: #e8f5e9;
                border-left: 5px solid #4CAF50;
            }
            .advice-header {
                color: #2E7D32;
                font-size: 20px;
                margin-bottom: 15px;
            }
            </style>
        """, unsafe_allow_html=True)

        st.header("🤖 Your Personalized Financial Advice")

        # Amount and Goals first
        col1, col2 = st.columns(2)
        
        with col1:
            amount = st.number_input(
                "Enter the amount you want to invest (₹)",
                min_value=500,
                value=20000,
                step=1000
            )

        with col2:
            horizons = st.multiselect(
                "Select your investment horizon",
                options=["Short Term", "Medium Term", "Long Term"],
                default=["Long Term"]
            )

        # Investment Types and Allocation
        types = st.multiselect("Choose preferred investment instruments", 
            ["SIP", "Mutual Fund", "Bond", "Gold", "Crypto"])

        # Show allocation section if types are selected
        allocation = {}
        show_allocation = False  # Initialize outside the if block
        if types:
            show_allocation = st.checkbox("Want to specify allocation percentages?")
            if show_allocation:
                col1, col2 = st.columns(2)
                
                with col1:
                    st.markdown("#### 📊 Allocate your investment (%)")
                    total = 0
                    equal_split = int(100 / len(types))
                    
                    # Display allocation inputs in rows
                    for inv_type in types:
                        percent = st.slider(
                            f"{inv_type}",
                            min_value=0,
                            max_value=100,
                            value=equal_split,
                            key=f"alloc_{inv_type}",
                            format="%d%%"  # Show percentage symbol
                        )
                        allocation[inv_type] = percent
                        total += percent
                    
                    if total != 100:
                        st.warning(f"Total allocation is {total}%. It should be 100%")
                    else:
                        st.success("Perfect! Allocation equals 100%")
                
                with col2:
                    st.markdown("#### 📈 Allocation Visualization")
                    # Create pie chart using plotly
                    fig = go.Figure(data=[go.Pie(
                        labels=list(allocation.keys()),
                        values=list(allocation.values()),
                        hole=.3,
                        textinfo='label+percent'
                    )])
                    fig.update_layout(
                        showlegend=False,
                        margin=dict(t=0, b=0, l=0, r=0),
                        height=300
                    )
                    st.plotly_chart(fig, use_container_width=True)

        # Submit button in a form
        with st.form("financial_advisor_form", clear_on_submit=False):
            submitted = st.form_submit_button(
                "Get Financial Advice",
                use_container_width=True
            )

            if submitted:
                # Create investment preferences data
                investment_preferences = {
                    "amount": amount,
                    "types": types if types else [], 
                    "horizons": horizons if types else [],
                    "allocation": allocation if show_allocation and types else None
                }
                st.session_state.investment_preferences = investment_preferences
                st.session_state.form_submitted = True

        # Display advice if form is submitted
        if st.session_state.get("form_submitted", False):
            with st.spinner('Generating your personalized financial advice...'):
                preferences_data = st.session_state.investment_preferences
                response = financial_adviser(preferences_data)
                
                # Display financial advice in a structured format
                st.markdown("### 💰 Financial Advice")

                # Create tabs for different aspects of advice
                advice_tab1, advice_tab2 = st.tabs([
                    "💡 Key Recommendations",
                    "📊 Investment Strategy",
                ])

                with advice_tab1:
                    st.markdown(response)
                    st.markdown("</div>", unsafe_allow_html=True)

                with advice_tab2:
                    st.header("Investment Strategy Visualizations")
                    
                    try:
                        # Get visualization suggestions from the AI
                        chart_data = investment_strategy_visulizations(response)
                        
                        if not chart_data:
                            st.info("No visualization data could be extracted from the advice.")
                            return
                            
                        if not chart_data.get("suggested_charts"):
                            st.info("No charts were suggested for the current investment strategy.")
                            return
                            
                        for chart in chart_data["suggested_charts"]:
                            try:
                                # Validate chart data
                                required_fields = ["type", "title", "x_axis", "y_axis"]
                                missing_fields = [field for field in required_fields if field not in chart]
                                
                                if missing_fields:
                                    st.warning(f"Skipping chart due to missing fields: {', '.join(missing_fields)}")
                                    continue
                                    
                                if not chart["x_axis"] or not chart["y_axis"]:
                                    st.warning(f"Skipping '{chart['title']}' due to empty data")
                                    continue
                                
                                # Create and display chart
                                with st.expander(f"📊 {chart['title']}", expanded=True):
                                    try:
                                        fig = plot_suggested_charts({"suggested_charts": [chart]})
                                        if fig:
                                            st.plotly_chart(fig, use_container_width=True)
                                            st.markdown(f"**Description:** {chart.get('description', 'No description available')}")
                                        else:
                                            st.warning(f"Could not create visualization for '{chart['title']}'")
                                    except ValueError as ve:
                                        st.error(str(ve))
                                        st.info("Please check the data format in the AI response.")
                                    except Exception as ce:
                                        st.error(f"Error creating chart: {str(ce)}")
                                        
                            except Exception as chart_error:
                                st.error(f"Error processing chart '{chart.get('title', 'Unnamed')}': {str(chart_error)}")
                                continue  # Continue with next chart even if one fails
                            
                    except json.JSONDecodeError:
                        st.error("Error parsing visualization data. The AI response was not in the expected format.")
                        st.info("Try adjusting your investment preferences or try again.")
                    except Exception as viz_error:
                        st.error(f"Error creating visualizations: {str(viz_error)}")
                        st.info("Unable to generate charts for this advice. Please try adjusting your investment preferences.")

                # Add download button for detailed report
                st.download_button(
                    label="📥 Download Detailed Financial Report",
                    data=str(response),
                    file_name="financial_advice_report.txt",
                    mime="text/plain"
                )

    # File uploader in the sidebar
    uploaded_file = st.sidebar.file_uploader("Upload PDF, DOCX or CSV files", type=["pdf", "docx", "csv"])

    # Disable the button if no files are uploaded
    if uploaded_file:
        upload_button_disabled = False
    else:
        upload_button_disabled = True

    if st.sidebar.button("Upload & Process Files", disabled=upload_button_disabled):
        if uploaded_file:
            try:
                # Process the file
                st.session_state.df, st.session_state.text_content = vector_embedding(uploaded_file)
                st.session_state.file_processed = True
                st.sidebar.success(
                    "Files processed! You can now ask questions about your financial documents or manage your budget."
                )
            except Exception as e:
                st.sidebar.error(f"Error processing file: {str(e)}")
                st.session_state.file_processed = False
                st.session_state.df = None
                st.session_state.text_content = None
        else:
            st.sidebar.error("Please upload files first.")

    with tab1:
        try:
            # Initialize container for messages
            messages_container = st.container()
            input_placeholder = st.empty()

            # Display messages in the container
            with messages_container:
                try:
                    for message in st.session_state.conversation_history:
                        with st.chat_message(message["role"], avatar="🧑‍💼" if message["role"] == "user" else "🤖"):
                            st.write(message["content"])
                            if message["role"] == "assistant" and "response_time" in message:
                                st.markdown(
                                    f'<div class="response-time">Response time: {message["response_time"]:.2f}s</div>',
                                    unsafe_allow_html=True)
                except Exception as e:
                    st.error(f"Error displaying messages: {str(e)}")
                    # Reset conversation history if corrupted
                    st.session_state.conversation_history = []

            # Place the input at the bottom
            with input_placeholder:
                prompt = st.chat_input("Ask me about your financial documents...", key="chat_input")

            if prompt:
                try:
                    # Validate input
                    if not prompt.strip():
                        st.warning("Please enter a valid question.")
                        return

                    # Check if documents are processed
                    if "vectors" not in st.session_state:
                        st.error("Please upload and process documents first! Use the sidebar to upload your files.")
                        return

                    # Add user message to history
                    st.session_state.conversation_history.append({
                        "role": "user",
                        "content": prompt,
                        "timestamp": time.time()
                    })

                    # Process message and show response
                    with st.chat_message("assistant", avatar="🤖"):
                        with st.spinner("Processing your question..."):
                            try:
                                # Initialize chain components
                                document_chain = create_stuff_documents_chain(llm, prompt_temp)
                                retriever = st.session_state.vectors.as_retriever()
                                retrieval_chain = create_retrieval_chain(retriever, document_chain)

                                # Process query with timeout
                                start_time = time.time()
                                try:
                                    response = retrieval_chain.invoke({'input': prompt})
                                    response_time = time.time() - start_time

                                    if not response or 'answer' not in response:
                                        raise ValueError("No valid response received from the model")

                                    # Add to conversation history
                                    st.session_state.conversation_history.append({
                                        "role": "assistant",
                                        "content": response['answer'],
                                        "response_time": response_time,
                                        "timestamp": time.time()
                                    })

                                except Exception as e:
                                    st.error(f"Error processing your question: {str(e)}")
                                    st.info("Try rephrasing your question or uploading different documents.")
                                    # Remove the failed user message
                                    if st.session_state.conversation_history:
                                        st.session_state.conversation_history.pop()

                            except Exception as chain_error:
                                st.error(f"Error setting up the processing chain: {str(chain_error)}")
                                st.info(
                                    "There might be an issue with the document processing. Try uploading your documents again.")
                                if st.session_state.conversation_history:
                                    st.session_state.conversation_history.pop()

                    # Refresh the chat display
                    st.rerun()

                except Exception as e:
                    st.error(f"An unexpected error occurred: {str(e)}")
                    st.info("Please try again or refresh the page if the problem persists.")
                    # Clean up any partial state
                    if st.session_state.conversation_history:
                        st.session_state.conversation_history.pop()

        except Exception as e:
            st.error("Critical error in chat interface. Please refresh the page.")
            st.error(f"Error details: {str(e)}")
            # Reset the session state
            st.session_state.conversation_history = []

    with tab2:
        try:
            st.header("Financial Dashboard")

            if not uploaded_file:
                st.info("👋 Welcome to the Financial Dashboard! Please upload a file using the sidebar to get started.")
                st.markdown("""
                    #### Supported File Types:
                    - 📊 CSV files (for structured financial data)
                    - 📄 PDF documents
                    - 📝 Word documents (DOCX)

                    Upload your financial documents to see automatic analysis and visualizations!
                """)
                return

            try:
                # Get the current dataframe from session state
                df = st.session_state.df
                text_content = st.session_state.text_content

                if df is not None:
                    st.session_state.data = df

                    # Try to analyze with LLM
                    try:
                        analysis = analyze_data_with_llm(df, text_content)
                        if analysis:
                            st.session_state.analysis_results = analysis
                            st.session_state.chart_config = analysis.get("suggested_charts", [])
                        else:
                            st.warning("Could not generate automatic analysis. Showing raw data instead.")
                    except Exception as analysis_error:
                        st.error(f"Error during data analysis: {str(analysis_error)}")
                        st.info("Showing raw data visualization options instead of AI-powered analysis.")
                        analysis = None

                    # Data summary section
                    with st.expander("Data Summary", expanded=True):
                        try:
                            # Display basic info
                            col1, col2, col3 = st.columns(3)
                            with col1:
                                st.metric("Rows", f"{len(df)}")
                            with col2:
                                st.metric("Columns", f"{len(df.columns)}")
                            with col3:
                                st.metric("Data Types", f"{len(df.dtypes.unique())}")

                            # Display LLM summary if available
                            if analysis and "summary" in analysis:
                                st.markdown(f"**Summary:** {analysis['summary']}")

                            # Display insights if available
                            if analysis and "insights" in analysis and analysis["insights"]:
                                st.markdown("**Key Insights:**")
                                for insight in analysis["insights"]:
                                    st.markdown(f"- {insight}")

                            # Data preview with error handling
                            tab1, tab2 = st.tabs(["Data Preview", "Column Details"])
                            with tab1:
                                try:
                                    st.dataframe(df.head(10), use_container_width=True)
                                except Exception as preview_error:
                                    st.error(f"Error displaying data preview: {str(preview_error)}")

                            with tab2:
                                try:
                                    col_info = pd.DataFrame({
                                        "Column": df.columns,
                                        "Type": df.dtypes,
                                        "Non-Null Count": df.count(),
                                        "Null Count": df.isna().sum(),
                                        "Unique Values": [df[col].nunique() for col in df.columns]
                                    })
                                    st.dataframe(col_info, use_container_width=True)
                                except Exception as col_info_error:
                                    st.error(f"Error displaying column details: {str(col_info_error)}")

                        except Exception as summary_error:
                            st.error(f"Error generating data summary: {str(summary_error)}")

                    # Visualizations section
                    st.header("Auto-Generated Visualizations")

                    try:
                        # If we have charts to display
                        if st.session_state.chart_config:
                            # Create a grid for charts
                            charts_per_row = 2
                            num_charts = len(st.session_state.chart_config)
                            rows = (num_charts + charts_per_row - 1) // charts_per_row

                            for row in range(rows):
                                cols = st.columns(charts_per_row)
                                for col_idx in range(charts_per_row):
                                    chart_idx = row * charts_per_row + col_idx
                                    if chart_idx < num_charts:
                                        with cols[col_idx]:
                                            try:
                                                chart_cfg = st.session_state.chart_config[chart_idx]
                                                fig = create_visualization(chart_cfg, df)
                                                if fig:
                                                    st.plotly_chart(fig, use_container_width=True)
                                                    with st.expander("Description"):
                                                        st.write(
                                                            chart_cfg.get("description", "No description available"))
                                                else:
                                                    st.warning("Could not create this visualization")
                                            except Exception as chart_error:
                                                st.error(f"Error creating chart {chart_idx + 1}: {str(chart_error)}")
                        else:
                            st.info("No automatic visualizations available. Try the custom visualization tools below.")

                    except Exception as viz_error:
                        st.error(f"Error in visualization section: {str(viz_error)}")

                    # Custom visualization section with error handling
                    with st.expander("Create Custom Visualization"):
                        try:
                            col1, col2 = st.columns(2)
                            with col1:
                                chart_type = st.selectbox(
                                    "Chart Type",
                                    ["line", "bar", "scatter", "pie", "box", "histogram"]
                                )
                            with col2:
                                chart_title = st.text_input("Chart Title", "My Custom Chart")

                            col1, col2, col3 = st.columns(3)
                            with col1:
                                x_axis = st.selectbox("X-Axis", ["None"] + list(df.columns))
                                x_axis = None if x_axis == "None" else x_axis
                            with col2:
                                y_axis = st.selectbox("Y-Axis", ["None"] + list(df.columns))
                                y_axis = None if y_axis == "None" else y_axis
                            with col3:
                                color_by = st.selectbox("Color By", ["None"] + list(df.columns))
                                color_by = None if color_by == "None" else color_by

                            if st.button("Generate Custom Chart"):
                                try:
                                    if not x_axis and not y_axis:
                                        st.warning("Please select at least one axis for the chart.")
                                        return

                                    custom_config = {
                                        "type": chart_type,
                                        "title": chart_title,
                                        "x_axis": x_axis,
                                        "y_axis": y_axis,
                                        "color": color_by
                                    }

                                    custom_fig = create_visualization(custom_config, df)
                                    if custom_fig:
                                        st.plotly_chart(custom_fig, use_container_width=True)
                                    else:
                                        st.warning("Could not create visualization with the selected parameters")
                                except Exception as custom_chart_error:
                                    st.error(f"Error creating custom chart: {str(custom_chart_error)}")
                                    st.info("Try different parameters or check if the selected columns are compatible.")

                        except Exception as custom_viz_error:
                            st.error(f"Error in custom visualization section: {str(custom_viz_error)}")

                else:
                    if not st.session_state.file_processed:
                        st.warning(
                            "Please process the uploaded file using the 'Upload & Process Files' button in the sidebar.")
                    else:
                        st.error("No valid data found in the uploaded file. Please check the file format and contents.")

            except Exception as data_error:
                st.error(f"Error processing data: {str(data_error)}")
                st.info("Try uploading a different file or check if the file format is correct.")

        except Exception as tab_error:
            st.error("Critical error in dashboard. Please refresh the page and try again.")
            st.error(f"Error details: {str(tab_error)}")

        finally:
            # Show raw document text if available
            if 'raw_text' in st.session_state and st.session_state.raw_text:
                with st.expander("Document Text"):
                    try:
                        st.text_area("Extracted Text", st.session_state.raw_text, height=200)
                    except Exception as text_error:
                        st.error(f"Error displaying extracted text: {str(text_error)}")

    # CSS for styling
    st.markdown("""
        <style>
        /* Main container adjustments */
        .main .block-container {
            padding-bottom: 100px !important;
        }

        /* Stacked container layout */
        .stChatFloatingInputContainer {
            position: fixed !important;
            bottom: 0 !important;
            left: 50% !important;
            transform: translateX(-50%) !important;
            width: 100% !important;
            max-width: 800px !important;
            padding: 20px !important;
            background: linear-gradient(to top, rgba(255,255,255,1) 50%, rgba(255,255,255,0.9) 100%) !important;
            backdrop-filter: blur(10px) !important;
            border-top: 1px solid rgba(0,0,0,0.1) !important;
            z-index: 9999 !important;
        }

        /* Message container adjustments */
        .stChatMessageContent {
            max-width: 800px !important;
            margin: 0 auto !important;
        }

        /* Ensure messages don't get hidden behind input */
        .stChatMessageContainer {
            padding-bottom: 100px !important;
        }

        /* Custom scrollbar for WebKit browsers */
        ::-webkit-scrollbar {
            width: 8px;
        }

        ::-webkit-scrollbar-track {
            background: transparent;
        }

        ::-webkit-scrollbar-thumb {
            background: rgba(0,0,0,0.2);
            border-radius: 4px;
        }

        ::-webkit-scrollbar-thumb:hover {
            background: rgba(0,0,0,0.3);
        }
        </style>
    """, unsafe_allow_html=True)

    # Add a sidebar instructions box
    st.sidebar.info("""    
    ℹ️ **How it works**

    \n This app automatically creates dashboards from your documents:

    1. **Upload a file** - CSV, PDF, or Word document
    2. **Process the document** - We'll extract data and text
    3. **AI Analysis** - Using Groq or Gemini to identify patterns
    4. **Auto Visualizations** - Generate charts based on data structure
    5. **Customize** - Create your own visualizations
    5. **Financial Planning** - Recommendation based on income, expanse and future goals

    For best results:
    - CSV files work best for structured data analysis
    - Make sure your PDFs or Word docs have table-like data for visualization
    """)

    st.sidebar.markdown('© 2025 Financial Assistant AI.')

    # Cleanup temporary directory on session state initialization
    def cleanup_temp_dir():
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)
            os.makedirs(temp_dir)

    if "initialized" not in st.session_state:
        st.session_state.initialized = True
        cleanup_temp_dir()


if __name__ == "__main__":
    main_document()